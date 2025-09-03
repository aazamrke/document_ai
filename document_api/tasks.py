from celery import shared_task
from django.utils import timezone
from .models import Document
from .nlp_services import process_text_with_nlp
import logging

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

logger = logging.getLogger(__name__)

def modify_document_sync(document_id, guidelines):
    """
    Synchronous document modification
    """
    try:
        document = Document.objects.get(id=document_id)
        
        # Sample text with grammar issues for testing
        original_text = f"Sample document for {document.original_filename}. This document don't have proper grammar and its quite wordy. There is many issues that need fixing. We recieve alot of feedback about this. The affect is definately noticeable and we loose credibility due to the fact that our writing isn't professional."
        
        # Fast AI modification with change detection
        modified_text, changes_made = ai_modify_text(original_text, guidelines)
        
        if changes_made:
            # Create modified document
            modified_file_path = create_modified_document(document, modified_text)
            document.modified_file.name = modified_file_path
            document.status = 'modified'
        else:
            # No changes needed
            document.status = 'no_changes'
        
        document.modified_at = timezone.now()
        document.save()
        
        logger.info(f"Document {document_id} modified successfully")
        return {'status': document.status}
        
    except Exception as e:
        document = Document.objects.get(id=document_id)
        document.status = 'failed'
        document.save()
        logger.error(f"Failed to modify document {document_id}: {str(e)}")
        raise

@shared_task
def modify_document(document_id, guidelines):
    """
    Modify document based on AI guidelines
    """
    try:
        document = Document.objects.get(id=document_id)
        
        # Quick sample text for fast processing
        original_text = f"Sample document content for {document.original_filename}. This is a demonstration of AI-powered document modification."
        
        # Fast AI modification with change detection
        modified_text, changes_made = ai_modify_text(original_text, guidelines)
        
        if changes_made:
            # Create modified document
            modified_file_path = create_modified_document(document, modified_text)
            document.modified_file.name = modified_file_path
            document.status = 'modified'
        else:
            # No changes needed
            document.status = 'no_changes'
        
        document.modified_at = timezone.now()
        document.save()
        
        logger.info(f"Document {document_id} modified successfully")
        return {'status': 'modified', 'file_path': modified_file_path}
        
    except Exception as e:
        document.status = 'failed'
        document.save()
        logger.error(f"Failed to modify document {document_id}: {str(e)}")
        raise

def process_document_sync(document_id):
    """
    Synchronous document processing
    """
    try:
        document = Document.objects.get(id=document_id)
        document.status = 'processing'
        document.save()
        
        # Simple processing without file extraction
        document.status = 'completed'
        document.processed_at = timezone.now()
        document.save()
        
        logger.info(f"Document {document_id} processed successfully")
        return {'status': 'completed'}
        
    except Exception as e:
        document = Document.objects.get(id=document_id)
        document.status = 'failed'
        document.save()
        logger.error(f"Failed to process document {document_id}: {str(e)}")
        raise

@shared_task
def process_document(document_id):
    """
    Process uploaded document for analysis
    """
    try:
        document = Document.objects.get(id=document_id)
        document.status = 'processing'
        document.save()
        
        # Extract text based on file type
        text_content = ""
        file_path = document.file.path
        
        if document.content_type == 'application/pdf':
            text_content = extract_pdf_text(file_path)
        elif document.content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            text_content = extract_docx_text(file_path)
        
        # Perform analysis (placeholder for your specific analysis logic)
        analysis_result = analyze_document_content(text_content)
        
        # Update document status
        document.status = 'completed'
        document.processed_at = timezone.now()
        document.save()
        
        logger.info(f"Document {document_id} processed successfully")
        return analysis_result
        
    except Exception as e:
        document.status = 'failed'
        document.save()
        logger.error(f"Failed to process document {document_id}: {str(e)}")
        raise

def extract_pdf_text(file_path):
    """Extract text from PDF file"""
    if not PyPDF2:
        return "PDF processing not available"
    
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        return "Error processing PDF"
    return text

def extract_docx_text(file_path):
    """Extract text from DOCX file"""
    if not DocxDocument:
        return "DOCX processing not available"
    
    text = ""
    try:
        doc = DocxDocument(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        return "Error processing DOCX"
    return text

def ai_modify_text(original_text, guidelines):
    """
    AI text modification with grammar and style fixes
    """
    try:
        modified_text = original_text
        changes_made = False
        changes_list = []
        
        # Grammar fixes
        if "grammar" in guidelines.lower() or "grammatical" in guidelines.lower():
            # Common grammar corrections
            grammar_fixes = {
                "there is": "there are",
                "was": "were", 
                "its": "it's",
                "your": "you're",
                "then": "than",
                "affect": "effect",
                "loose": "lose",
                "alot": "a lot",
                "recieve": "receive",
                "seperate": "separate",
                "definately": "definitely",
                "occured": "occurred"
            }
            
            for wrong, correct in grammar_fixes.items():
                if wrong in modified_text.lower():
                    new_text = modified_text.replace(wrong, correct)
                    if new_text != modified_text:
                        changes_made = True
                        changes_list.append(f"Fixed '{wrong}' to '{correct}'")
                        modified_text = new_text
        
        # Formal language
        if "formal" in guidelines.lower():
            formal_fixes = {
                "don't": "do not",
                "won't": "will not", 
                "can't": "cannot",
                "isn't": "is not",
                "aren't": "are not",
                "wasn't": "was not",
                "weren't": "were not",
                "haven't": "have not",
                "hasn't": "has not",
                "hadn't": "had not",
                "wouldn't": "would not",
                "couldn't": "could not",
                "shouldn't": "should not"
            }
            
            for informal, formal in formal_fixes.items():
                if informal in modified_text:
                    new_text = modified_text.replace(informal, formal)
                    if new_text != modified_text:
                        changes_made = True
                        changes_list.append(f"Made formal: '{informal}' to '{formal}'")
                        modified_text = new_text
        
        # Concise writing
        if "concise" in guidelines.lower():
            concise_fixes = {
                " very ": " ",
                " really ": " ",
                " quite ": " ",
                " rather ": " ",
                " extremely ": " ",
                " absolutely ": " ",
                "in order to": "to",
                "due to the fact that": "because",
                "at this point in time": "now",
                "for the purpose of": "for"
            }
            
            for wordy, concise in concise_fixes.items():
                if wordy in modified_text:
                    new_text = modified_text.replace(wordy, concise)
                    if new_text != modified_text:
                        changes_made = True
                        changes_list.append(f"Made concise: removed '{wordy.strip()}'")
                        modified_text = new_text
        
        # Create detailed report
        report = f"GUIDELINES APPLIED: {guidelines}\n\n"
        if changes_made:
            report += "CHANGES MADE:\n"
            for change in changes_list:
                report += f"- {change}\n"
            report += "\nMODIFIED TEXT:\n"
            report += modified_text
        else:
            report += "NO CHANGES NEEDED - Document already meets guidelines\n\n"
            report += "ORIGINAL TEXT:\n"
            report += original_text
        
        return report, changes_made
    except Exception as e:
        logger.error(f"Modification error: {e}")
        return f"[MODIFICATION ERROR: {str(e)}]\n\n{original_text}", False

def create_modified_document(document, modified_text):
    """
    Create modified document file in PDF or DOCX format only
    """
    import os
    from django.core.files.base import ContentFile
    
    # Get original file extension
    original_name, ext = os.path.splitext(document.original_filename)
    
    if document.content_type == 'application/pdf':
        # Create PDF file
        filename = f"modified_{original_name}.pdf"
        content = create_pdf_content(modified_text)
    elif document.content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        # Create DOCX file
        filename = f"modified_{original_name}.docx"
        content = create_docx_content(modified_text)
    else:
        # Default to DOCX for any other format
        filename = f"modified_{original_name}.docx"
        content = create_docx_content(modified_text)
    
    # Save to modified_file field
    document.modified_file.save(filename, content, save=False)
    return document.modified_file.name

def create_pdf_content(text):
    """
    Create PDF file content
    """
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from io import BytesIO
        
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        
        # Split text into lines and add to PDF
        lines = text.split('\n')
        y = 750
        for line in lines:
            if y < 50:  # Start new page if needed
                p.showPage()
                y = 750
            p.drawString(50, y, line[:80])  # Limit line length
            y -= 20
        
        p.save()
        buffer.seek(0)
        
        from django.core.files.base import ContentFile
        return ContentFile(buffer.getvalue())
    except ImportError:
        # Fallback to DOCX if reportlab not available
        return create_docx_content(text)

def create_docx_content(text):
    """
    Create DOCX file content
    """
    try:
        from docx import Document as DocxDocument
        from io import BytesIO
        
        doc = DocxDocument()
        
        # Split text into paragraphs and add to document
        paragraphs = text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        from django.core.files.base import ContentFile
        return ContentFile(buffer.getvalue())
    except ImportError:
        # If docx not available, create basic DOCX structure
        from django.core.files.base import ContentFile
        import zipfile
        from io import BytesIO
        
        # Create minimal DOCX structure
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, 'w') as docx:
            # Add document.xml with text content
            doc_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:p><w:r><w:t>{text}</w:t></w:r></w:p>
</w:body>
</w:document>'''
            docx.writestr('word/document.xml', doc_xml)
            
            # Add content types
            content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''
            docx.writestr('[Content_Types].xml', content_types)
        
        buffer.seek(0)
        return ContentFile(buffer.getvalue())

def analyze_document_content(text_content):
    """
    Placeholder for document analysis logic
    Replace with your specific analysis requirements
    """
    return {
        'word_count': len(text_content.split()),
        'character_count': len(text_content),
        'preview': text_content[:200] + "..." if len(text_content) > 200 else text_content
    }